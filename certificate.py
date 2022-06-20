from docx import *
from docx.blkcntnr import BlockItemContainer
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.section import Section, Sections
from docx.shared import *
docPath = "C:\\Users\\Ibrahim.Moghul\\Desktop\\Data Analysis Scripts\\CSV OUTPUT\\TONGRUN\\Certificates\\"
docName = "test.docx"
xmlName = "test.xml"
doc = Document(docPath+docName)
# print(document.tables)
from docx import Document

print(doc.tables)
exit()

# doc.tables[-1].cell(0, 0).text = "stuff"
# doc.tables[-1].cell(0, 0).width = 500
# doc.tables[-1].cell(0, 1).text = "stuff"

for row in doc.tables[0].rows:
    for cell in row.cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.size= Pt(30)
                font.name = "AvenirNext LT Pro Regular"
                font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
doc.save(docPath+docName)
