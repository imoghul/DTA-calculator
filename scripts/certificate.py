from docx import *
from docx.shared import *
import contextlib
from docx2pdf import convert
import os
import sys
import shutil
import glob

def createDocPath(sn, cbDate, path): # convert sn, date, and path into a path for the certificate
    name = path + sn + "_" + cbDate.replace("/", "") + "_certificate"
    num = len(glob.glob(name+"*"))
    name = name+("(%d)"%num if num else "") +".docx" 
    return name

# def getDocPath(sn, cbDate, path):
#     name = path + sn + "_" + cbDate.replace("/", "") + "_certificate"
#     return glob.glob(name+"*").sort(reverse=True)[0]

def createCopy(sn, cbDate, path): # create a copy of TEMPLATE.docx to a new path for the specified certificate
    dest = createDocPath(sn, cbDate, path)
    shutil.copy2(path + "TEMPLATE.docx", dest)
    return dest


def createCertificate(sn, cbDate, result, DAQTemp, PostCalibAir, path, header = False):
    try:
        dest = createCopy(sn, cbDate, path) # create the copy
        # dest = getDocPath(sn, cbDate, path)
        doc = Document(dest) # open the copy

        for t in doc.tables: # loop through the tables, becuase whoever make the certificate put the data we need to change in invisible tables
            for c in t._cells: # loop through cells
                if c.text == "SERIAL NUMBER": # change SERIAL NUMBER to match desired sn and match font
                    c.text = sn
                    paragraphs = c.paragraphs
                    paragraph = paragraphs[0]
                    run_obj = paragraph.runs
                    run = run_obj[0]
                    font = run.font
                    font.size = Pt(30)
                    font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                    font.name = "AvenirNext LT Pro Regular"
                elif (
                    c.text == "CALIBRATION DATE"
                    or c.text == "RESULT"
                    or c.text == "DAQ TEMP"
                    or c.text == "CALIB"
                    or c.text == "Air Monitoring Probe:"
                ): # set other cells
                # NOTE: for some reason, separating this if into multiple ones messed with the font 
                    try: # insert data into desired places
                        if c.text == "CALIBRATION DATE":
                            c.text = cbDate
                        elif c.text == "RESULT":
                            c.text = result
                        elif c.text == "DAQ TEMP":
                            c.text = DAQTemp
                        elif c.text == "CALIB":
                            c.text = PostCalibAir
                        elif c.text == "Air Monitoring Probe:" and header:
                            c.text = "Air Monitoring Probe:\nGlycol Probe:" # if header is true, add a LF and Glycol Probe:
                    except:
                        raise Exception(
                            "Error while generating the certificate")
                    paragraphs = c.paragraphs
                    paragraph = paragraphs[0]
                    run_obj = paragraph.runs
                    run = run_obj[0]
                    font = run.font
                    font.size = Pt(12)
                    font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                    font.name = "AvenirNext LT Pro Regular"

                doc.save(dest)
    except Exception as e:
        try: # remove the copy if a failure occured during the certificate generation, to avoid half generated certificates
            os.remove(dest)#(getDocPath(sn, cbDate, path))
        except:pass
        raise e


def convertToPDF_doc(doc): # convert a single docx to a pdf
    # with contextlib.redirect_stdout(open(os.devnull, 'w')):
        try:
            print("Converting %s to pdf" % doc.replace("\\", "/").split("/")[-1])
            convert(doc)
        except:
            print(doc + " couldn't be converted to a pdf")


def convertToPDF_path(path): # convert all docx in path to pdf
    # with contextlib.redirect_stdout(open(os.devnull, 'w')):
        try:
            print("Converting certificates to pdfs")
            convert(path)
        except Exception as e:
            print("Some files couldn't be converted to a pdf")
